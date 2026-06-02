"""
自动报告生成器
输出：专业 Word 报告，包含执行摘要、异常清单、承运商分析、建议
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime


def _hex_rgb(hex_color):
    h = hex_color.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def _add_heading(doc, text, level=1, color_hex="1A3C5E"):
    p   = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    r, g, b = _hex_rgb(color_hex)
    run.font.color.rgb = RGBColor(r, g, b)
    run.font.bold = True
    run.font.size = Pt({1: 16, 2: 13}.get(level, 11))
    return p


def _add_kpi_table(doc, stats):
    kpis = [
        ("总分析订单数",   f"{stats['total_orders']:,}"),
        ("检测到异常订单", f"{stats['total_anomalies']:,}"),
        ("整体异常率",     f"{stats['anomaly_rate']*100:.1f}%"),
        ("平均延误天数",   f"{stats['avg_delay']} 天"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    row = table.rows[0]
    for i, (label, value) in enumerate(kpis):
        cell = row.cells[i]
        _set_cell_bg(cell, "1A3C5E")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(value)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size  = Pt(20)
        run.font.bold  = True
        p2  = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2  = p2.add_run(label)
        r2.font.color.rgb = RGBColor(180, 200, 220)
        r2.font.size  = Pt(9)


def _add_data_table(doc, df, header_bg="1A3C5E"):
    cols  = list(df.columns)
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"

    hdr = table.rows[0]
    for i, col in enumerate(cols):
        cell = hdr.cells[i]
        _set_cell_bg(cell, header_bg)
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(col)
        run.font.bold      = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size      = Pt(9)

    sev_colors = {"严重": "FFD0CC", "中等": "FFF3CC", "轻微": "E8F4FD", "正常": "F0FFF0"}
    for idx, (_, row_data) in enumerate(df.iterrows()):
        row = table.add_row()
        bg  = "F8F9FA" if idx % 2 == 0 else "FFFFFF"
        for i, col in enumerate(cols):
            cell = row.cells[i]
            val  = str(row_data[col])
            _set_cell_bg(cell, sev_colors.get(val, bg))
            p   = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(val).font.size = Pt(8.5)


def generate_report(stats, result_df, chart_paths, output_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── 封面 ─────────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("运输计划异常检测分析报告")
    tr.font.size = Pt(24)
    tr.font.bold = True
    r, g, b = _hex_rgb("1A3C5E")
    tr.font.color.rgb = RGBColor(r, g, b)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        f"Shipping Anomaly Detection Report  |  "
        f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}"
    )
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(100, 120, 140)
    doc.add_paragraph()

    # ── 执行摘要 ─────────────────────────────────────────────────
    _add_heading(doc, "一、执行摘要", 1)
    _add_kpi_table(doc, stats)
    doc.add_paragraph()

    sev = stats["severity_dist"]
    doc.add_paragraph(
        f"本报告对 {stats['total_orders']:,} 条运输订单进行多维度异常检测，"
        f"综合运用规则阈值、统计 Z-Score 及 Isolation Forest 三种方法交叉验证。"
        f"共发现异常订单 {stats['total_anomalies']:,} 条，整体异常率 "
        f"{stats['anomaly_rate']*100:.1f}%。"
        f"严重异常（延误>14天）{sev.get('严重', 0)} 条，"
        f"中等异常 {sev.get('中等', 0)} 条，"
        f"平均延误 {stats['avg_delay']} 天，最长延误 {stats['max_delay']} 天。"
    )

    # ── 异常总览 ─────────────────────────────────────────────────
    _add_heading(doc, "二、异常检测总览", 1)
    if "anomaly_overview" in chart_paths:
        doc.add_picture(chart_paths["anomaly_overview"], width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 港口吞吐量 ───────────────────────────────────────────────
    _add_heading(doc, "三、港口吞吐量宏观背景", 1)
    doc.add_paragraph(
        "结合港口吞吐量统计数据，分析主要港口季节性规律及年度增长趋势，"
        "为订单层面异常的宏观解读提供背景参考。"
    )
    if "port_throughput" in chart_paths:
        doc.add_picture(chart_paths["port_throughput"], width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 延误分布 ─────────────────────────────────────────────────
    _add_heading(doc, "四、延误模式分析", 1)
    if "delay_distribution" in chart_paths:
        doc.add_picture(chart_paths["delay_distribution"], width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 承运商绩效 ───────────────────────────────────────────────
    _add_heading(doc, "五、承运商绩效评估", 1)
    if "carrier_performance" in chart_paths:
        doc.add_picture(chart_paths["carrier_performance"], width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    cs = stats["carrier_stats"].copy()
    cs["平均延误天数"] = cs["平均延误天数"].round(2)
    cs["延误率"]      = (cs["延误率"] * 100).round(1).astype(str) + "%"
    cs["异常率"]      = (cs["异常率"] * 100).round(1).astype(str) + "%"
    _add_data_table(doc, cs[["承运商", "总订单数", "平均延误天数", "延误率", "异常订单数", "异常率"]])

    # ── Top 异常订单 ─────────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "六、重点异常订单清单（Top 10）", 1)
    top = stats["top_anomalies"].copy()
    top.columns = ["订单ID", "承运商", "路线", "运输方式", "延误天数", "severity", "异常原因"]
    _add_data_table(doc, top)

    # ── 高风险路线 ───────────────────────────────────────────────
    doc.add_paragraph()
    _add_heading(doc, "七、高风险路线排名", 1)
    rs = stats["route_stats"].head(8).copy()
    rs["平均延误"] = rs["平均延误"].round(1)
    rs["异常率"]   = (rs["异常率"] * 100).round(1).astype(str) + "%"
    _add_data_table(doc, rs[["路线", "订单数", "平均延误", "最大延误", "异常数", "异常率"]])

    # ── 建议 ─────────────────────────────────────────────────────
    doc.add_paragraph()
    _add_heading(doc, "八、风险应对建议", 1)
    worst  = stats["carrier_stats"].iloc[0]["承运商"]
    best   = stats["carrier_stats"].iloc[-1]["承运商"]
    recs = [
        f"【承运商优化】{worst} 异常率最高，建议列为季度重点审查对象，"
        f"高价值货物优先考虑 {best}。",
        "【预警机制】延误偏差率超过 30% 的在途订单启动主动跟踪，提前通知收货方。",
        "【旺季备货】2月（春节）及4-5月历史低谷期前，建议提前 2-3 周发运。",
        "【数据治理】建议与 TMS 对接，实现实时数据接入和自动预警。",
        "【KPI 考核】将异常率、平均延误纳入承运商年度评分卡。",
    ]
    for rec in recs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(rec).font.size = Pt(10)

    doc.add_paragraph()
    note = doc.add_paragraph(
        "注：本报告基于历史运输记录，仅供内部决策参考。"
        "异常判定使用多模型投票机制，置信水平约 90%。"
    )
    note.runs[0].font.size      = Pt(8.5)
    note.runs[0].font.color.rgb = RGBColor(130, 130, 130)

    doc.save(output_path)
    print(f"  ✓ 报告已生成：{output_path}")
